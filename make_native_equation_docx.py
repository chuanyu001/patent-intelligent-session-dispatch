from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUT_DIR = next((ROOT / "outputs").iterdir())
SRC = next(OUT_DIR.glob("*20260626141330.md"))
DST = OUT_DIR / f"{SRC.stem}_公式简化原生版.docx"


INLINE_MATH_RE = re.compile(
    r"(\\\(((?:\\.|[^)])+?)\\\)|(?<!\$)\$(?!\$)((?:\\.|[^$\n])+?)\$(?!\$))"
)


def set_font(run, name="宋体", size=10.5, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def latex_to_linear_math(text: str) -> str:
    s = text.strip()
    s = re.sub(r"^\\\(|\\\)$", "", s).strip()
    s = re.sub(r"^\\\[|\\\]$", "", s).strip()
    s = re.sub(r"^\$+|\$+$", "", s).strip()
    s = s.replace("\\left", "").replace("\\right", "")
    s = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", s)
    s = re.sub(r"\\operatorname\{([^{}]+)\}", r"\1", s)
    replacements = {
        r"\Longleftrightarrow": "⇔",
        r"\Longleftarrow": "⇐",
        r"\Rightarrow": "⇒",
        r"\geq": "≥",
        r"\ge": "≥",
        r"\leq": "≤",
        r"\le": "≤",
        r"\neq": "≠",
        r"\land": "∧",
        r"\wedge": "∧",
        r"\lor": "∨",
        r"\vee": "∨",
        r"\exists": "∃",
        r"\forall": "∀",
        r"\in": "∈",
        r"\varnothing": "∅",
        r"\ldots": "…",
        r"\Delta": "Δ",
        r"\theta": "θ",
        r"\lambda": "λ",
        r"\tau": "τ",
        r"\quad": "    ",
        r"\,": " ",
        r"\;": " ",
        r"\\": "\n",
    }
    for k, v in replacements.items():
        s = s.replace(k, v)
    s = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", s)
    s = s.replace("\\ ", " ")
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n\s+", "\n", s)
    return s.strip()


def add_math(paragraph, latex: str):
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = latex_to_linear_math(latex)
    mr.append(mt)
    omath.append(mr)
    paragraph._p.append(omath)


def add_inline(paragraph, text: str):
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_font(run)
        latex = m.group(2) or m.group(3) or ""
        run = paragraph.add_run(latex_to_linear_math(latex))
        set_font(run)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run)


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    add_inline(p, text)


def split_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def add_table(doc: Document, rows: list[str]):
    data = [split_table_row(r) for r in rows if not is_separator_row(r)]
    if not data:
        return
    cols = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(data):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, row[j] if j < len(row) else "")
            for run in p.runs:
                set_font(run, bold=(i == 0))


def add_image(doc: Document, src: str):
    path = (OUT_DIR / src).resolve()
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.5))


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    lines = SRC.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("<!-- !["):
            m = re.search(r"\]\(([^)]+)\)", stripped)
            if m and "math_figures" not in m.group(1):
                add_image(doc, m.group(1))
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code))
            set_font(run, "Consolas", 9)
            i += 1
            continue
        if stripped.startswith("\\[") or stripped.startswith("$$"):
            end_marker = "\\]" if stripped.startswith("\\[") else "$$"
            math_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(end_marker):
                math_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_math(p, "\n".join(math_lines))
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            h = doc.add_heading(m.group(2), level=level)
            for run in h.runs:
                set_font(run, "黑体" if level <= 2 else "宋体")
            i += 1
            continue
        add_paragraph(doc, stripped)
        i += 1

    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    build()
