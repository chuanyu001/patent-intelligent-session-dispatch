from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "智能会话派单运营监控"
FIG_DIR = OUT_DIR / "visio_figures_expanded"
STAMP = datetime.now().strftime("%Y%m%d%H%M%S")
CASE = "一种企业即时通信消息流的会话切分、事件聚合及任务派发方法"
BASE = f"一种面向企业即时通信客户群的智能会话派单与运营监控方法、系统、设备及存储介质_{STAMP}_扩展规范版"
DOCX_PATH = OUT_DIR / f"{BASE}.docx"
MD_PATH = OUT_DIR / f"{BASE}.md"
PLAN_PATH = OUT_DIR / f"项目可扩展专利方案清单_{STAMP}_扩展规范版.md"
LOG_PATH = OUT_DIR / "交底书修订对话记录.md"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, width: int, fnt) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            trial = line + ch
            if draw.textlength(trial, font=fnt) <= width:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = ch
        lines.append(line)
    return lines


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#5B6B7C", width=3):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        pts = [(x2, y2), (x2 - sign * 12, y2 - 7), (x2 - sign * 12, y2 + 7)]
    else:
        sign = 1 if dy >= 0 else -1
        pts = [(x2, y2), (x2 - 7, y2 - sign * 12), (x2 + 7, y2 - sign * 12)]
    draw.polygon(pts, fill=color)


def poly_arrow(draw: ImageDraw.ImageDraw, points, color="#5B6B7C", width=3):
    draw.line(points, fill=color, width=width)
    arrow(draw, points[-2], points[-1], color, width)


def box(draw: ImageDraw.ImageDraw, xy, title: str, body: str = "", fill="#FFFFFF"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline="#7E8EA0", width=2)
    tf = font(23, True)
    bf = font(18)
    draw.text((x1 + 16, y1 + 12), title, fill="#17202A", font=tf)
    if body:
        for i, line in enumerate(wrap_text(draw, body, x2 - x1 - 32, bf)[:5]):
            draw.text((x1 + 16, y1 + 48 + i * 25), line, fill="#3B4A5A", font=bf)


def make_figures():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    specs = [
        ("01_system_architecture.jpg", make_system_architecture),
        ("02_method_flow.jpg", make_method_flow),
        ("03_dual_threshold.jpg", make_dual_threshold),
        ("04_session_state.jpg", make_session_state),
        ("05_dedup_event.jpg", make_dedup_event),
        ("06_dispatch_decision.jpg", make_dispatch_decision),
        ("07_alarm_state.jpg", make_alarm_state),
        ("08_feedback_tuning.jpg", make_feedback_tuning),
        ("09_data_objects.jpg", make_data_objects),
    ]
    for name, fn in specs:
        fn(FIG_DIR / name)


def make_system_architecture(path: Path):
    img = Image.new("RGB", (1600, 820), "#F7F9FB")
    d = ImageDraw.Draw(img)
    nodes = [
        ("接入层", "消息采集\n身份归并\n员工/机器人过滤", "#E9F3FF"),
        ("状态处理层", "SessionBuffer\n双阈值闭合\n分布式锁与版本", "#EFF7EE"),
        ("语义处理层", "格式归一\n意图识别\n事件指纹", "#FFF7E8"),
        ("任务决策层", "重复抑制\n事件聚合\n候选人派发", "#FDEEEE"),
        ("监控反馈层", "SLA监听\n告警升级\n复核与调参", "#F1EEFF"),
    ]
    coords = []
    for i, (title, body, fill) in enumerate(nodes):
        x = 55 + i * 305
        xy = (x, 110, x + 250, 330)
        coords.append(xy)
        box(d, xy, title, body, fill)
    for a, b in zip(coords, coords[1:]):
        arrow(d, (a[2], 220), (b[0], 220))
    lower = [
        ((210, 505, 510, 710), "策略配置库", "阈值、候选规则\n值班表、权限、版本"),
        ((650, 505, 950, 710), "任务与工单库", "会话、任务、父事件\n子会话、告警、复核"),
        ((1090, 505, 1390, 710), "运营看板", "漏单、重复、准确率\n延迟、回滚记录"),
    ]
    for xy, title, body in lower:
        box(d, xy, title, body, "#FFFFFF")
    arrow(d, (360, 505), (360, 330))
    arrow(d, (800, 330), (800, 505))
    arrow(d, (1090, 610), (950, 610))
    poly_arrow(d, [(1240, 505), (1240, 435), (340, 435), (340, 505)])
    img.save(path, quality=95)


def make_method_flow(path: Path):
    img = Image.new("RGB", (1600, 1000), "#FFFFFF")
    d = ImageDraw.Draw(img)
    rows = [
        [("S1", "采集消息"), ("S2", "身份归并"), ("S3", "写入会话缓存"), ("S4", "双阈值闭合")],
        [("S8", "同客重复抑制"), ("S7", "事件指纹/相似度"), ("S6", "意图识别"), ("S5", "格式归一")],
        [("S9", "跨客事件聚合"), ("S10", "候选处理人生成"), ("S11", "派发并去重发送"), ("S12", "SLA监听")],
        [("S16", "策略版本发布"), ("S15", "参数评估"), ("S14", "人工复核"), ("S13", "告警升级")],
    ]
    coords = {}
    for r, row in enumerate(rows):
        for c, (code, label) in enumerate(row):
            x = 70 + c * 380
            y = 60 + r * 230
            coords[code] = (x, y, x + 285, y + 135)
            box(d, coords[code], code, label, "#F8FBFF" if r < 2 else "#FAFAFA")
    order = ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "S9", "S10", "S11", "S12", "S13", "S14", "S15", "S16"]
    for left, right in zip(order, order[1:]):
        a, b = coords[left], coords[right]
        if b[1] == a[1]:
            if b[0] > a[0]:
                arrow(d, (a[2], (a[1] + a[3]) // 2), (b[0], (b[1] + b[3]) // 2))
            else:
                arrow(d, (a[0], (a[1] + a[3]) // 2), (b[2], (b[1] + b[3]) // 2))
        else:
            arrow(d, ((a[0] + a[2]) // 2, a[3]), ((b[0] + b[2]) // 2, b[1]))
    img.save(path, quality=95)


def make_dual_threshold(path: Path):
    img = Image.new("RGB", (1500, 760), "#FFFFFF")
    d = ImageDraw.Draw(img)
    y = 360
    d.line([(120, y), (1360, y)], fill="#5B6B7C", width=4)
    points = [(180, "m1\n09:00"), (340, "m2\n09:01"), (520, "m3\n09:03"), (900, "m4\n09:09")]
    for x, label in points:
        d.ellipse((x - 16, y - 16, x + 16, y + 16), fill="#2D7DD2")
        d.text((x - 38, y + 28), label, font=font(19), fill="#17202A")
    d.arc((520, 230, 900, 490), 185, 355, fill="#D9534F", width=4)
    arrow(d, (885, 340), (900, y), "#D9534F", 3)
    d.text((620, 235), "Δt3 ≥ T_gap", font=font(24, True), fill="#D9534F")
    d.arc((180, 120, 760, 610), 185, 355, fill="#7E57C2", width=4)
    arrow(d, (745, 330), (760, y), "#7E57C2", 3)
    d.text((345, 120), "L(g,u,k) ≥ T_len", font=font(24, True), fill="#7E57C2")
    box(d, (1050, 120, 1380, 265), "闭合条件", "空闲定时器触发\n或最长持续定时器触发", "#FFF7E8")
    box(d, (1050, 475, 1380, 620), "输出", "标准化会话任务\n写入版本号与消息列表", "#EFF7EE")
    arrow(d, (1215, 265), (1215, 475))
    img.save(path, quality=95)


def make_session_state(path: Path):
    img = Image.new("RGB", (1500, 760), "#FFFFFF")
    d = ImageDraw.Draw(img)
    nodes = [
        ((95, 275, 340, 430), "空闲", "无打开缓存"),
        ((470, 275, 715, 430), "打开", "追加消息\n刷新idle_timer"),
        ((845, 275, 1090, 430), "待闭合", "获取分布式锁\n校验版本"),
        ((1215, 275, 1460, 430), "已闭合", "生成任务\n释放缓存"),
    ]
    for xy, title, body in nodes:
        box(d, xy, title, body, "#F8FBFF")
    for a, b in zip(nodes, nodes[1:]):
        arrow(d, (a[0][2], 352), (b[0][0], 352))
    d.text((370, 230), "首条客户消息", font=font(20), fill="#3B4A5A")
    d.text((742, 230), "T_gap 或 T_len", font=font(20), fill="#3B4A5A")
    d.text((1115, 230), "锁成功", font=font(20), fill="#3B4A5A")
    poly_arrow(d, [(965, 430), (965, 565), (590, 565), (590, 430)])
    d.text((690, 575), "锁失败或版本变化：回到打开态", font=font(20), fill="#D9534F")
    img.save(path, quality=95)


def make_dedup_event(path: Path):
    img = Image.new("RGB", (1500, 860), "#FFFFFF")
    d = ImageDraw.Draw(img)
    box(d, (120, 110, 520, 285), "同客户重复抑制", "同一群、同一客户\nT_dup窗口内相似诉求\n追加到既有工单", "#EFF7EE")
    box(d, (120, 500, 520, 675), "跨客户事件聚合", "不同客户、同一事件指纹\n相似度超过θ_event\n父事件+子会话", "#FFF7E8")
    box(d, (780, 110, 1160, 285), "已有工单", "ticket_id\nmessage_id_list\nSLA更新", "#FFFFFF")
    box(d, (780, 500, 1160, 675), "父事件任务", "event_id\nassociated_customer_ids\nchild_session_ids", "#FFFFFF")
    arrow(d, (520, 197), (780, 197))
    arrow(d, (520, 587), (780, 587))
    d.line([(675, 70), (675, 725)], fill="#C9D1D9", width=3)
    d.text((605, 742), "两种机制独立判定，不互相替代", font=font(24, True), fill="#17202A")
    img.save(path, quality=95)


def make_dispatch_decision(path: Path):
    img = Image.new("RGB", (1500, 920), "#FFFFFF")
    d = ImageDraw.Draw(img)
    box(d, (90, 90, 370, 245), "输入任务", "群类型、意图\n置信度、SLA等级", "#E9F3FF")
    box(d, (520, 90, 800, 245), "候选生成", "服务组织\n角色/技能/权限", "#EFF7EE")
    box(d, (950, 90, 1230, 245), "约束过滤", "排班、在线状态\n负载上限、权限", "#FFF7E8")
    box(d, (520, 380, 800, 535), "优先级排序", "时段目标\n意图目标\n默认目标", "#FDEEEE")
    box(d, (950, 380, 1230, 535), "发送控制", "幂等键\n失败重试\n重复发送抑制", "#FFFFFF")
    box(d, (520, 670, 800, 825), "兜底/转派", "无候选\n离线\n超负载", "#F1EEFF")
    arrow(d, (370, 167), (520, 167))
    arrow(d, (800, 167), (950, 167))
    arrow(d, (1090, 245), (1090, 380))
    arrow(d, (950, 457), (800, 457))
    arrow(d, (660, 535), (660, 670))
    poly_arrow(d, [(800, 747), (1090, 747), (1090, 535)])
    img.save(path, quality=95)


def make_alarm_state(path: Path):
    img = Image.new("RGB", (1500, 760), "#FFFFFF")
    d = ImageDraw.Draw(img)
    nodes = [
        ((90, 280, 330, 430), "待响应", "R(x)=0"),
        ((450, 280, 690, 430), "轻度超时", "T_light触发\n提醒处理人"),
        ((810, 280, 1050, 430), "严重超时", "T_severe触发\n升级负责人"),
        ((1170, 280, 1410, 430), "已响应", "R(x)=1\n关闭监听"),
    ]
    for xy, title, body in nodes:
        box(d, xy, title, body, "#F8FBFF")
    arrow(d, (330, 355), (450, 355))
    arrow(d, (690, 355), (810, 355))
    arrow(d, (1050, 355), (1170, 355))
    poly_arrow(d, [(570, 280), (570, 180), (1290, 180), (1290, 280)])
    poly_arrow(d, [(930, 430), (930, 560), (1290, 560), (1290, 430)])
    d.text((665, 150), "任一阶段收到有效回复", font=font(20), fill="#3B4A5A")
    d.text((980, 575), "高风险意图可直接进入严重告警", font=font(20), fill="#D9534F")
    img.save(path, quality=95)


def make_feedback_tuning(path: Path):
    img = Image.new("RGB", (1500, 820), "#FFFFFF")
    d = ImageDraw.Draw(img)
    nodes = [
        ((95, 130, 345, 285), "采样窗口", "任务、工单\n回复、复核"),
        ((470, 130, 720, 285), "指标计算", "MissRate\nDupRate\nAccuracy\nDelayRate"),
        ((845, 130, 1095, 285), "参数评估", "候选阈值组合\n目标函数J"),
        ((1220, 130, 1470, 285), "策略发布", "版本、生效范围\n发布时间"),
        ((845, 520, 1095, 675), "人工审批", "高风险策略\n灰度发布"),
        ((470, 520, 720, 675), "回滚", "指标恶化\n恢复上一版本"),
    ]
    for xy, title, body in nodes:
        box(d, xy, title, body, "#F8FBFF")
    for a, b in zip(nodes[:4], nodes[1:4]):
        arrow(d, (a[0][2], 207), (b[0][0], 207))
    arrow(d, (970, 285), (970, 520))
    arrow(d, (1220, 207), (1095, 595))
    arrow(d, (845, 595), (720, 595))
    poly_arrow(d, [(470, 595), (220, 595), (220, 285)])
    img.save(path, quality=95)


def make_data_objects(path: Path):
    img = Image.new("RGB", (1500, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    objects = [
        ((70, 90, 390, 245), "MessageEvent", "message_id\ngroup_id sender_id\nmessage_type send_time"),
        ((590, 90, 910, 245), "SessionBuffer", "session_id\ngroup_id customer_id\nmessage_id_list timers"),
        ((1110, 90, 1430, 245), "Task", "task_id source_session\nintent confidence\nassigned_target"),
        ((590, 560, 910, 715), "Ticket/Event", "ticket_id/event_id\nparent-child\nassociated_customer_ids"),
        ((1110, 560, 1430, 715), "ReviewRecord", "expected_boundary\nexpected_assignee\nduplicate/missed flag"),
    ]
    for xy, title, body in objects:
        box(d, xy, title, body, "#FFFFFF")
    arrow(d, (390, 167), (590, 167))
    arrow(d, (910, 167), (1110, 167))
    arrow(d, (1270, 245), (750, 560))
    arrow(d, (1270, 245), (1270, 560))
    arrow(d, (1110, 637), (910, 637))
    img.save(path, quality=95)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run("（" + str(items.index(item) + 1) + "）" + item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def math_run(text: str):
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def math_sub(base: str, sub: str):
    node = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(math_run(base))
    s = OxmlElement("m:sub")
    s.append(math_run(sub))
    node.append(e)
    node.append(s)
    return node


def math_frac(num: list, den: list):
    f = OxmlElement("m:f")
    num_node = OxmlElement("m:num")
    den_node = OxmlElement("m:den")
    append_math_parts(num_node, num)
    append_math_parts(den_node, den)
    f.append(num_node)
    f.append(den_node)
    return f


def append_math_parts(parent, parts: list):
    for part in parts:
        if isinstance(part, str):
            parent.append(math_run(part))
        elif part[0] == "sub":
            parent.append(math_sub(part[1], part[2]))
        elif part[0] == "frac":
            parent.append(math_frac(part[1], part[2]))


def add_formula(doc: Document, parts: list):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    math = OxmlElement("m:oMath")
    append_math_parts(math, parts)
    math_para = OxmlElement("m:oMathPara")
    math_para.append(math)
    p._p.append(math_para)


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if i == 0:
                r.bold = True
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / filename), width=Inches(6.4))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    add_heading(doc, "技术交底书", 1)
    add_para(doc, f"建议技术名称：{CASE}", True)
    add_para(doc, "专利类型：发明。技术联系人、电话、邮箱：待补充。", True)

    add_heading(doc, "一、相似技术交底书写法参考与本稿优化原则", 1)
    add_para(doc, "相似工单、会话处理和智能客服类专利文本通常采用“技术领域、背景技术、要解决的技术问题、技术方案、有益效果、附图说明、具体实施方式”的结构。在技术方案部分，较成熟的写法不是只描述业务流程，而是把消息流、状态缓存、任务对象、判定条件、阈值、版本发布和异常处理写成可被本领域技术人员复现的数据处理链路。")
    add_para(doc, "本次优化据此将原稿的运营描述改写为企业即时通信消息流处理方案，重点突出异步消息接入、按群和客户维度的状态化聚合、空闲与最长持续双定时器闭合、同客户重复抑制、跨客户事件聚合、候选处理人约束过滤、响应监听以及复核反馈调参。")
    add_table(doc, [
        ["参考写法要点", "本稿对应优化"],
        ["问题必须落在技术对象上", "将漏单、重复派单改写为消息流乱序、状态缓存、计时闭合、事件关联和幂等发送问题"],
        ["公式符号先定义后使用", "集中给出Δt、L、T_gap、T_duration、θ_dup、θ_event、R(x)等定义"],
        ["实施例要有时间线和数据对象", "新增同客户连续消息、同客户重复、跨客户公共事件和异常场景实施例"],
        ["附图服务于技术链路", "新增9张Visio风格图，图内不放标题，连线尽量横平竖直且不交叉"],
    ])

    add_heading(doc, "二、技术领域", 1)
    add_para(doc, "本发明涉及企业即时通信、客户群消息处理、智能工单派发和运营监控技术领域，尤其涉及一种对企业即时通信客户群中的异步多模态消息进行会话切分、事件聚合、任务派发、响应监听和质量反馈的方法、系统、电子设备及计算机可读存储介质。")

    add_heading(doc, "三、背景技术与现有技术缺陷", 1)
    add_para(doc, "企业服务通常通过企业微信、钉钉、飞书或类似即时通信平台维护大量客户群。客户在群内以文字、语音、图片、链接、文件、小程序卡片等形式提出咨询、报修、投诉、配件、权益和预约等诉求。由于消息天然呈异步流式到达，同一客户会分多条消息补充上下文，不同客户也可能在同一时间讨论同一公共事件。")
    add_para(doc, "如果系统逐条消息判断是否建单，容易造成上下文缺失；如果只按固定时间窗口聚合，又会在客户持续补充信息时长期不闭合。现有智能客服或工单系统通常偏重问答、摘要、单轮意图识别或人工分配提醒，缺少面向客户群消息流的状态化会话缓存、双阈值切分、同客户重复抑制、跨客户事件聚合、候选处理人约束派发和响应质量闭环的组合方案。")
    add_table(doc, [
        ["公开资料", "可借鉴内容", "与本方案差异"],
        ["CN116032871B，基于聊天消息创建工单的方法、装置和设备", "从客户群聊天消息判断是否创建工单", "未公开同群同客户双阈值会话闭合、同客户重复抑制、跨客户事件聚合及反馈调参链路"],
        ["CN112686674A，一种客服对话工单总结方法，申请日2020-12-25，公开日2021-04-20，申请人科讯嘉联信息技术有限公司", "根据对话主题和槽位提取模型生成工单总结", "重点是对话主题拆解和槽位提取，不处理客户群异步消息流的切分、派发、告警和版本化策略更新"],
        ["企业即时通信/SCRM/会话存档类公开产品", "提供消息接入、客户关系和会话留痕能力", "通常未披露可专利化的状态缓存、定时器闭合、事件指纹聚合和派发监控闭环"],
        ["国外自动客服、意图识别、对话摘要类方案", "可用于语义识别或工单摘要", "不针对中国企业即时通信客户群的多客户交叉发言、群类型路由和运营监控场景"],
    ])
    add_para(doc, "因此，本方案的创新重点不放在单纯使用大模型或语义模型识别意图，而放在消息流状态管理、会话任务边界、重复与聚合的分离判定、候选处理人生成、监控开关解耦以及基于复核指标的策略版本更新。")

    add_heading(doc, "四、要解决的技术问题", 1)
    add_bullets(doc, [
        "在异步客户群消息流中，如何按群和客户维度维持会话状态，使多条补充消息形成完整任务，同时避免任务长期不闭合。",
        "如何在分布式接入、乱序到达和重复回调情况下避免同一会话被重复闭合或重复发送工单。",
        "如何区分同一客户在短时间内重复提出的同一诉求，与多个客户围绕同一公共事件提出的相似诉求。",
        "如何根据群类型、意图、技能标签、值班时段、在线状态、权限和负载上限生成处理人候选，并在无候选或发送失败时进行兜底。",
        "如何在不必对所有消息都派单的情况下保持响应监听、超时告警和高风险意图升级。",
        "如何将人工复核结果转化为可执行的阈值、相似度、兜底策略和发布版本调整。"
    ])

    add_heading(doc, "五、系统总体结构", 1)
    add_figure(doc, "01_system_architecture.jpg", "图1 系统总体结构示意图")
    add_para(doc, "系统包括接入层、状态处理层、语义处理层、任务决策层、监控反馈层以及数据存储。接入层采集平台消息回调并进行身份归并；状态处理层维护SessionBuffer并执行双阈值闭合；语义处理层进行格式归一、意图识别和事件指纹生成；任务决策层执行重复抑制、事件聚合和候选处理人派发；监控反馈层负责响应监听、告警升级、人工复核和策略调参。")

    add_heading(doc, "六、整体方法流程", 1)
    add_figure(doc, "02_method_flow.jpg", "图2 整体方法流程示意图")
    add_bullets(doc, [
        "接收客户群消息事件，记录平台标识、消息标识、群标识、发送者标识、发送者类型、消息类型、发送时间和接收时间。",
        "通过成员表、托管账号表、机器人表、客户身份映射表进行身份归并，过滤员工、机器人、系统通知和黑名单消息。",
        "以(group_id, customer_id)为键写入SessionBuffer，追加message_id_list，刷新last_message_time并维护idle_timer和duration_timer。",
        "当空闲阈值或最长持续阈值满足时，获取分布式锁并校验缓存版本，闭合会话并生成标准化任务。",
        "对任务内容进行格式归一、意图识别、事件指纹生成，并分别进入同客户重复抑制和跨客户事件聚合。",
        "未被合并的任务根据候选处理人生成、约束过滤和多级优先级排序完成派发，并写入幂等发送键。",
        "系统独立建立响应监听，根据轻度超时、严重超时和高风险意图触发告警。",
        "人工复核样本形成质量指标，策略评估模块在有界候选参数内选择新版本并支持灰度发布和回滚。"
    ])

    add_heading(doc, "七、核心数据对象", 1)
    add_figure(doc, "09_data_objects.jpg", "图3 核心数据对象关联示意图")
    add_table(doc, [
        ["对象", "关键字段"],
        ["MessageEvent", "message_id、platform_id、group_id、sender_id、sender_type、message_type、content、media_reference、send_time、receive_time、reply_reference"],
        ["SessionBuffer", "session_id、group_id、customer_id、message_id_list、first_message_time、last_message_time、session_status、idle_timer、duration_timer、version"],
        ["Task", "task_id、source_session_id、normalized_content、intent_category、intent_confidence、dispatchable_flag、event_fingerprint、associated_customer_ids、assigned_target、task_status、sla_level、creation_time"],
        ["ReviewRecord", "review_id、task_id、expected_boundary、expected_intent、expected_assignee、duplicate_flag、missed_flag、reviewer、review_time"],
        ["StrategyVersion", "version_id、scope、T_gap、T_duration、T_dup、T_event、θ_dup、θ_event、publish_time、rollback_version、approval_status"],
    ])

    add_heading(doc, "八、符号、参数与公式", 1)
    add_table(doc, [
        ["符号", "含义"],
        ["g、u、k", "群标识、客户标识、同一(group_id, customer_id)下的消息序号"],
        ["Δt_k", "第k条消息与第k-1条消息的发送或接收时间差"],
        ["L(g,u,k)", "同一群同一客户当前会话从首条消息到第k条消息的持续时长"],
        ["T_gap、T_duration", "空闲闭合阈值、最长持续时长阈值"],
        ["x、y", "新任务和已存在工单或事件任务"],
        ["g_x、g_y、u_x、u_y、Δt(x,y)", "任务x、y的群、客户以及任务间时间差"],
        ["θ_dup、θ_event", "同客户重复抑制相似度阈值、跨客户事件聚合相似度阈值"],
        ["N_user、A_amb、S_avg", "事件窗口内客户数、公共事件歧义标志、平均相似度"],
        ["R(x)、t(x,first)", "任务x是否已有有效回复、任务x首次待响应时间"],
        ["N_correct、N_ticket、N_miss、N_need、N_dup", "正确任务数、任务总数、漏单数、应建单需求数、重复任务数"],
    ])
    add_formula(doc, [("sub", "Close", "g,u,k"), "=1 ⇔ (", ("sub", "Δt", "k"), "≥", ("sub", "T", "gap"), ") ∨ (", ("sub", "L", "g,u,k"), "≥", ("sub", "T", "duration"), ")"])
    add_formula(doc, ["Merge(x,y)=1 ⇔ (", ("sub", "g", "x"), "=", ("sub", "g", "y"), ") ∧ (", ("sub", "u", "x"), "=", ("sub", "u", "y"), ") ∧ (S(x,y)≥", ("sub", "θ", "dup"), ") ∧ (", ("sub", "Δt", "x,y"), "≤", ("sub", "T", "dup"), ")"])
    add_formula(doc, ["Event(x,y)=1 ⇔ (", ("sub", "g", "x"), "=", ("sub", "g", "y"), ") ∧ (", ("sub", "u", "x"), "≠", ("sub", "u", "y"), ") ∧ (S(x,y)≥", ("sub", "θ", "event"), ") ∧ (", ("sub", "Δt", "x,y"), "≤", ("sub", "T", "event"), ")"])
    add_formula(doc, ["ParentEvent=1 ⇔ (", ("sub", "N", "user"), "≥2) ∧ (", ("sub", "A", "amb"), "=1) ∧ (", ("sub", "S", "avg"), "≥", ("sub", "θ", "event"), ") ∧ (", ("sub", "T", "duration"), "≤", ("sub", "T", "event"), ")"])
    add_formula(doc, ["D(x)=first(", ("sub", "D", "slot"), ",", ("sub", "D", "intent"), ",", ("sub", "D", "default"), ")"])
    add_formula(doc, ["Alarm(x)=1 ⇔ R(x)=0 ∧ now-", ("sub", "t", "x,first"), "≥", ("sub", "T", "light")])
    add_formula(doc, ["J=α·MissRate+β·DupRate+γ·(1-Accuracy)+δ·DelayRate"])
    add_formula(doc, ["Accuracy=", ("frac", [("sub", "N", "correct")], [("sub", "N", "ticket")]), "，MissRate=", ("frac", [("sub", "N", "miss")], [("sub", "N", "need")]), "，DupRate=", ("frac", [("sub", "N", "dup")], [("sub", "N", "ticket")])])
    add_para(doc, "上述first函数表示按顺序选取第一个非空且满足权限、在线、负载和幂等发送约束的目标集合。若目标集合包含多人，则继续按技能匹配度、当前负载、最近处理时间和稳定哈希排序选择；若目标集合为空，则进入兜底队列或负责人升级队列。")

    add_heading(doc, "九、双阈值会话切分机制", 1)
    add_figure(doc, "03_dual_threshold.jpg", "图4 双阈值会话切分示意图")
    add_figure(doc, "04_session_state.jpg", "图5 会话缓存状态机示意图")
    add_para(doc, "SessionBuffer以(group_id, customer_id)为主键建立。每收到一条客户消息，系统先判断消息是否已经处理，若message_id已存在则丢弃重复回调；若消息时间早于当前last_message_time但落入允许乱序窗口，则按send_time重新排序并更新版本；否则按到达顺序追加。")
    add_para(doc, "idle_timer在每次追加客户消息后重置，当当前时间与last_message_time之差达到T_gap时触发空闲闭合。duration_timer在首条消息进入缓存时启动，当当前时间与first_message_time之差达到T_duration时触发最长持续闭合。两个定时器任一触发时，闭合线程获取(group_id, customer_id, session_id)对应的分布式锁并校验version；校验成功后生成Task，校验失败则说明其他线程已处理或缓存已变化，当前闭合请求放弃。")

    add_heading(doc, "十、重复抑制与跨客户事件聚合", 1)
    add_figure(doc, "05_dedup_event.jpg", "图6 重复抑制与事件聚合示意图")
    add_table(doc, [
        ["机制", "对象范围", "判定条件", "输出结果"],
        ["同客户重复抑制", "同一group_id且同一customer_id", "T_dup窗口内意图相同或语义相似度不低于θ_dup", "不新建工单，将新消息追加到已有ticket并更新SLA"],
        ["跨客户事件聚合", "同一group_id但不同customer_id", "T_event窗口内事件指纹相同或平均相似度不低于θ_event，且N_user达到阈值", "生成或更新父事件任务，保留每个客户的子会话记录"],
    ])
    add_para(doc, "两种机制的目的不同：重复抑制避免同一客户反复催促造成重复工单；事件聚合避免公共事件被拆成多个无关联工单，但不会丢弃各客户的个体需求。父事件任务用于统一处理公共原因，子会话用于记录每个客户的具体诉求、附件、地址、权益状态和后续回复。")

    add_heading(doc, "十一、智能派发机制", 1)
    add_figure(doc, "06_dispatch_decision.jpg", "图7 派发决策示意图")
    add_para(doc, "智能派发不是简单采用固定处理人，而是先按服务组织、群类型、意图类别、角色标签、技能标签、值班表、在线状态、权限范围和负载上限生成候选集。候选集为空时，系统将任务派发到兜底队列、值班负责人或人工分拣池；候选人离线或发送失败时，按幂等键记录失败并重试，不重复向同一目标发送。")
    add_para(doc, "在不需要复杂评分模型的实施例中，系统可采用约束过滤加多级优先级：第一优先级为时段目标D_slot，第二优先级为意图目标D_intent，第三优先级为默认目标D_default。若多个候选均满足条件，则按技能匹配数量、当前未完成任务数、最近分配时间和稳定哈希排序确定最终处理人。")

    add_heading(doc, "十二、响应监听与告警升级", 1)
    add_figure(doc, "07_alarm_state.jpg", "图8 响应监听状态示意图")
    add_para(doc, "派发开关和监控开关相互解耦。对于dispatchable_flag为false但monitor_flag为true的任务，系统不生成外部工单，但仍监听群内是否出现员工有效回复。有效回复R(x)=1可以由员工身份、回复引用、@客户、语义关联和时间窗口共同判定。")
    add_para(doc, "当R(x)=0且等待时长达到T_light时，系统触发轻度超时提醒；达到T_severe时，系统触发严重超时并升级负责人。对投诉、退款、合同风险、舆情、监管投诉等高风险意图，即使未达到T_severe，也可直接触发意图告警。")

    add_heading(doc, "十三、质量反馈与策略版本更新", 1)
    add_figure(doc, "08_feedback_tuning.jpg", "图9 反馈调参与策略版本示意图")
    add_para(doc, "系统按日、周或预设样本窗口抽取任务、消息、工单和回复记录，由人工复核或抽样标注生成ReviewRecord。指标计算模块统计Accuracy、MissRate、DupRate和DelayRate，并在候选参数集合中评估目标函数J。候选参数包括T_gap、T_duration、T_dup、T_event、θ_dup、θ_event、轻重超时阈值、兜底目标和高风险意图列表。")
    add_para(doc, "策略版本发布时记录version_id、scope、publish_time、approval_status和rollback_version。对影响面较大的策略，系统先在指定群类型或指定服务组织内灰度发布，并比较发布前后指标。若J恶化超过阈值或人工确认误伤增加，则自动或人工回滚到上一版本。")

    add_heading(doc, "十四、具体实施例", 1)
    add_para(doc, "实施例一：同客户连续消息。客户A在群G1于09:00发送“空调不制冷”，09:01补充图片，09:03发送“昨晚开始的”。系统将三条消息写入SessionBuffer(G1,A)，message_id_list为[m1,m2,m3]。09:08达到T_gap=300秒，系统获取锁并生成任务x，normalized_content包含故障描述、图片引用和发生时间，intent_category为维修，intent_confidence为0.91。派发模块根据群类型为区域售后群、意图为维修、当前时段为工作日白班，选择维修坐席B并启动SLA监听。")
    add_para(doc, "实施例二：同客户重复诉求。客户A在09:20再次发送“还没人联系我，空调还是坏的”。系统生成新任务x2后，在T_dup=1800秒窗口内查到同客户已有维修工单y，且S(x2,y)=0.86≥θ_dup=0.78，于是不新建工单，而是将新消息追加到y，更新last_customer_message_time和SLA提醒等级。")
    add_para(doc, "实施例三：跨客户公共事件。客户A、客户C、客户D在同一区域群内于10分钟内分别发送“门店系统打不开”“收银后台登不上”“今天后台一直报错”。三条任务的客户不同，但事件指纹均指向系统登录异常，平均相似度0.82≥θ_event=0.75，系统生成父事件任务event_1，并保留三个子会话，避免公共事件多头处理，同时保留各门店的具体影响范围。")
    add_para(doc, "实施例四：异常处理。图片识别失败时，系统保留media_reference并将format_status置为media_pending；语音无法转写时，任务进入人工补录队列；意图置信度低于阈值时，进入兜底意图并标记review_required；无候选处理人时，进入兜底队列并通知负责人；群内只有机器人回复时不记为R(x)=1；非工作时间可触发自动安抚回复但继续监听人工响应；乱序或重复消息通过message_id、send_time和version校验处理。")

    add_heading(doc, "十五、测试数据与效果说明", 1)
    add_para(doc, "在一个测试样本中，选取若干客户群连续消息作为回放数据，比较人工处理、逐条消息建单和本方案三种方式。为避免绝对化表述，本文仅表述为“在该测试样本中未观察到漏单”，不宣称任何场景下漏单率恒为0。")
    add_table(doc, [
        ["指标", "人工/逐条基线", "本方案样本结果", "说明"],
        ["任务准确率", "依赖人工经验，波动较大", "样本中达到90%以上", "按人工复核正确任务数/任务总数统计"],
        ["漏单情况", "多条补充消息和图片消息易漏", "样本中未观察到漏单", "按应建单需求数和漏单数统计"],
        ["重复派发", "同客户催促可能重复建单", "重复任务显著减少", "同客户重复抑制窗口内合并"],
        ["多客户公共事件", "容易形成多个孤立工单", "形成父事件和子会话", "公共原因统一处理，个体需求保留"],
        ["响应延迟", "依赖人工盯群", "可触发轻重超时告警", "按首条待响应时间到有效回复时间统计"],
    ])

    add_heading(doc, "十六、保护重点与可扩展权利要求方向", 1)
    add_para(doc, "建议主权利要求收窄到如下连续技术链路：获取企业即时通信客户群消息；按群标识和客户标识维护状态化会话缓存；基于空闲阈值和最长持续阈值进行双阈值闭合；生成标准化会话任务；对同客户重复诉求执行抑制合并，对跨客户同事件执行事件聚合；根据候选处理人约束过滤和多级优先级执行派发；独立进行响应监听和质量反馈。")
    add_bullets(doc, [
        "从属方向一：SessionBuffer字段、定时器、分布式锁和版本校验。",
        "从属方向二：同客户重复抑制和跨客户事件聚合的不同窗口、相似度阈值和输出对象。",
        "从属方向三：群类型、意图、时段、技能、在线状态、权限和负载上限联合派发。",
        "从属方向四：派发开关与监控开关解耦、轻重超时和高风险意图告警。",
        "从属方向五：基于J函数的参数评估、策略版本、灰度发布和回滚。"
    ])

    add_heading(doc, "十七、可替代实施方式", 1)
    add_para(doc, "相似度S(x,y)可采用关键词规则、向量相似度、事件指纹匹配或多模型加权方式实现；意图识别可采用规则模型、分类模型、生成式模型或人工兜底方式实现；派发目标可以是个人、角色组、值班队列或外部工单系统。上述替代方式只要仍然执行状态化会话切分、重复/聚合分离、候选派发和监控反馈链路，均应落入本方案的技术构思范围。")

    doc.save(DOCX_PATH)


def write_markdown():
    md = dedent(f"""
    # 技术交底书扩展规范版

    建议技术名称：{CASE}

    本版按“相似技术交底书写法参考、技术领域、背景技术、技术问题、系统结构、方法流程、数据对象、公式、双阈值切分、重复抑制与事件聚合、智能派发、响应监听、质量反馈、实施例、测试数据、保护重点”重构。

    重点优化：
    - 将业务问题转化为异步消息流、状态缓存、双定时器闭合、幂等派发和监控反馈等技术问题。
    - 明确SessionBuffer、Task、ReviewRecord、StrategyVersion等数据对象。
    - 区分同客户重复抑制与跨客户事件聚合。
    - 增加候选处理人生成、约束过滤、多级优先级、兜底和重复发送抑制。
    - 增加质量目标函数J和策略版本发布/回滚机制。
    - 增加9张Visio风格附图，图内无标题。
    - 增加4个具体实施例和谨慎测试效果表述。
    """).strip()
    MD_PATH.write_text(md + "\n", encoding="utf-8")


def write_plan():
    PLAN_PATH.write_text(dedent("""
    # 项目可扩展专利方案清单

    1. 企业即时通信客户群消息流的双阈值会话切分方法。
    2. 面向客户群公共事件的跨客户事件指纹聚合方法。
    3. 基于群类型、意图、时段、技能和负载的候选处理人派发方法。
    4. 派发开关与监控开关解耦的客户群响应监听方法。
    5. 基于人工复核指标的会话切分与派发策略版本化调参方法。
    6. 多模态客户群消息的可派单任务归一化方法。
    7. 客户群同客户重复诉求的幂等合并与SLA更新方法。
    8. 面向客户群运营的父事件任务与子会话关联数据结构。
    """).strip() + "\n", encoding="utf-8")


def append_log():
    local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    utc = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    entry = dedent(f"""

    ## {local} / {utc}

    - 类型：按审查建议扩展优化
    - 用户说明摘要：按照建议强化技术问题、双阈值、重复/聚合、智能派发、质量反馈、实施例、公式规范和附图。
    - 本轮交付文件：{MD_PATH.name}；{DOCX_PATH.name}；{PLAN_PATH.name}
    - 修订摘要：新增9张Visio风格附图；扩展核心数据对象、符号公式、状态机、候选派发、响应监听、策略版本更新、实施例和测试表述；公式采用Word原生OMML对象生成。
    """)
    old = LOG_PATH.read_text(encoding="utf-8") if LOG_PATH.exists() else "# 交底书修订对话记录\n"
    LOG_PATH.write_text(old + entry, encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_figures()
    build_docx()
    write_markdown()
    write_plan()
    append_log()
    print(MD_PATH)
    print(DOCX_PATH)
    print(PLAN_PATH)


if __name__ == "__main__":
    main()
