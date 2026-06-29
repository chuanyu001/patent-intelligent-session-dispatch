from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "智能会话派单运营监控"
STAMP = datetime.now().strftime("%Y%m%d%H%M%S")
CASE = "一种面向企业即时通信客户群的智能会话派单与运营监控方法、系统、设备及存储介质"
BASE = f"{CASE}_{STAMP}"
MD_PATH = OUT_DIR / f"{BASE}.md"
DOCX_PATH = OUT_DIR / f"{BASE}.docx"
PLAN_PATH = OUT_DIR / f"项目可扩展专利方案清单_{STAMP}.md"
LOG_PATH = OUT_DIR / "交底书修订对话记录.md"
FIG_DIR = OUT_DIR / "visio_figures"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        p = Path(item)
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color="#5B6B7C", width=3):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        pts = [(x2, y2), (x2 - sign * 12, y2 - 7), (x2 - sign * 12, y2 + 7)]
    else:
        sign = 1 if dy >= 0 else -1
        pts = [(x2, y2), (x2 - 7, y2 - sign * 12), (x2 + 7, y2 - sign * 12)]
    draw.polygon(pts, fill=color)


def draw_polyline_arrow(draw: ImageDraw.ImageDraw, points, color="#5B6B7C", width=3):
    if len(points) < 2:
        return
    draw.line(points, fill=color, width=width, joint="curve")
    draw_arrow(draw, points[-2], points[-1], color=color, width=width)


def box(draw, xy, title, subtitle="", fill="#FFFFFF", outline="#7E8EA0"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2)
    title_font = font(22, True)
    body_font = font(17)
    draw.text((x1 + 16, y1 + 14), title, fill="#17202A", font=title_font)
    if subtitle:
        lines = subtitle.split("\n")
        for idx, line in enumerate(lines):
            draw.text((x1 + 16, y1 + 48 + idx * 24), line, fill="#3B4A5A", font=body_font)


def make_system_fig(path: Path):
    img = Image.new("RGB", (1600, 780), "#F7F9FB")
    d = ImageDraw.Draw(img)
    top = [
        ("接入层", "消息采集\n身份归并\n过滤名单", "#E9F3FF"),
        ("处理层", "会话缓存\n双阈值切分\n格式归一化\n意图识别", "#EFF7EE"),
        ("决策层", "重复抑制\n多人聚合\n群/意图/时段派单", "#FFF4E5"),
        ("监控层", "超时监听\n轻度/严重报警\n意图报警", "#FDEEEE"),
    ]
    xs = [80, 455, 830, 1205]
    boxes = []
    for x, (title, sub, fill) in zip(xs, top):
        xy = (x, 115, x + 280, 365)
        box(d, xy, title, sub, fill)
        boxes.append(xy)
    for a, b in zip(boxes, boxes[1:]):
        draw_arrow(d, (a[2], 240), (b[0], 240))
    feedback = (455, 510, 735, 705)
    data = (830, 510, 1110, 705)
    box(d, feedback, "反馈层", "人工复核\n指标统计\n策略调参", "#F1EEFF")
    box(d, data, "核心数据对象", "客户群消息事件\n会话任务 x\n工单 y\n报警事件", "#FFFFFF")
    draw_polyline_arrow(d, [(1345, 365), (1345, 470), (595, 470), (595, 510)])
    draw_arrow(d, (735, 608), (830, 608))
    draw_polyline_arrow(d, [(595, 510), (595, 430), (595, 365)])
    img.save(path)


def make_flow_fig(path: Path):
    img = Image.new("RGB", (1600, 980), "#FFFFFF")
    d = ImageDraw.Draw(img)
    steps = [
        ("S1", "采集群消息\n群ID/客户ID/时间/格式"),
        ("S2", "身份归并与过滤\n员工/机器人/黑名单"),
        ("S3", "会话聚合\n同群同客户缓存"),
        ("S4", "双阈值切分\nT_gap 或 T_len"),
        ("S5", "格式归一化\n语音/图片/链接等"),
        ("S6", "意图识别\n母意图/子意图"),
        ("S7", "重复抑制\nT_merge 内合并"),
        ("S8", "多人聚合\n相似度 theta"),
        ("S9", "联合派单\n时段优先"),
        ("S10", "超时/意图报警\n轻度/严重并行"),
        ("S11", "人工复核\n错派/漏派/多派"),
        ("S12", "策略调参\n反馈阈值与兜底策略"),
    ]
    coords = []
    for row in range(3):
        for col in range(4):
            i = row * 4 + col
            x = 70 + col * 380
            y = 70 + row * 280
            coords.append((x, y, x + 295, y + 165))
            code, text = steps[i]
            box(d, coords[-1], code, text, "#F8FBFF" if row < 2 else "#FAFAFA")
    order = [0, 1, 2, 3, 7, 6, 5, 4, 8, 9, 10, 11]
    for i in range(11):
        a = coords[order[i]]
        b = coords[order[i + 1]]
        if order[i] in (3, 4):
            draw_arrow(d, ((a[0] + a[2]) // 2, a[3]), ((b[0] + b[2]) // 2, b[1]))
        else:
            if b[0] > a[0]:
                draw_arrow(d, (a[2], (a[1] + a[3]) // 2), (b[0], (b[1] + b[3]) // 2))
            else:
                draw_arrow(d, (a[0], (a[1] + a[3]) // 2), (b[2], (b[1] + b[3]) // 2))
    d.text((415, 925), "反馈回路：按 Accuracy、MissRate、DupRate 调整切分、抑制、聚合与兜底派单策略", fill="#3B4A5A", font=font(20))
    img.save(path)


DISCLOSURE = dedent(
    r"""
    # 技术交底书

    **案件名称**：一种面向企业即时通信客户群的智能会话派单与运营监控方法、系统、设备及存储介质

    **技术联系人**：
    - 姓名：[待填写]
    - 电话：[待填写]
    - 邮箱：[待填写]

    **专利类型**：发明

    ---

    ## 注意事项

    （1）交底书应使代理人能看懂，尤其是背景技术和详细技术方案，一定要写得全面、清楚、完整；
    （2）技术的公开程度，应以本领域普通技术人员不需付出创造性劳动即可进行实施为准；
    （3）在与代理人沟通时，对于代理人咨询的技术问题，应给予回答并认真讲解，并且按要求及时正确地补充相应技术材料。

    ---

    ## 一、技术背景与现有技术

    ### 1.1 现有技术

    检索说明：本稿参考国家知识产权局专利公布公告系统、Google Patents、企业即时通信及智能客服公开资料，以“客户群消息”“群聊工单”“会话切分”“意图识别”“智能派单”“超时报警”“运营监控”等作为检索词，对已公开专利、产品资料和学术资料进行比对。正式申请前仍建议由代理机构按最终权利要求再做一次 CNIPA 检索。

    #### 方向一：企业即时通信与会话存档运营平台

    1. **企业微信（腾讯）**：企业微信提供企业通信、客户联系、开放接口和客户关系沉淀能力，是客户群通信的基础设施。其局限在于仅提供平台能力，未公开面向客户大群的“会话切分、意图派单、重复抑制、超时分级报警和质量闭环”的组合式调度方法。来源链接：https://www.tencent.com/zh-cn/business/wecom.html

    2. **微伴助手、微盛企微管家、尘锋 SCRM 等私域运营工具**：公开资料显示上述工具具备会话存档、客户运营、客服或工单相关能力，部分产品还提供 AI 识别或预警工单功能。其局限在于公开资料通常侧重私域营销、销售转化、客户画像或通用工单流转，未明确披露本方案的同群同客户双阈值会话任务生成、同一诉求抑制合并、多人相似事件聚合、群类型/意图/时段联合派单、轻重超时与意图报警并行、以及以准确率/漏单率/多派率为目标的闭环调参。来源链接：https://www.weibanzhushou.com/；https://www.wshoto.com/；https://www.dustess.com/

    #### 方向二：群聊消息到工单的直接转化

    3. **CN116032871B，基于聊天消息创建工单的方法、装置和设备**：该专利公开了获取客户群中的群消息数据、判断是否需要创建工单、生成工单并发送至工单系统的基础链路，是与本方案最接近的公开方案之一。其局限在于未公开以同群同客户为单位的相邻消息间隔阈值 \(T_{\mathrm{gap}}\) 与持续时长上限 \(T_{\mathrm{len}}\) 双阈值会话切分，也未公开重复抑制窗口 \(T_{\mathrm{merge}}\)、多人相似事件聚合、时段优先派单、轻重超时分级报警和质量闭环。来源链接：https://patents.google.com/patent/CN116032871B/zh

    4. **CN122226741A，基于群消息上下文感知的智能数据抽取与任务生成系统及方法**：该公开方案侧重对群消息进行上下文感知、数据抽取和任务生成，能够说明“群消息到任务”属于已有技术方向。其局限在于重点是上下文抽取和任务生成架构，未公开本方案的双阈值切分、抑制合并、多人聚合、派单和运营监控闭环。来源链接：https://patents.google.com/patent/CN122226741A/zh

    5. **CN121814728A，基于企业微信会话存档的客户会话分配与提醒方法及系统**：该公开方案涉及新增会话检测、基于排班和规则的会话分配以及超时提醒。其局限在于偏“分配 + 提醒”，未公开意图、群类型、时段和接收对象的联合决策，也未公开派单开关与监控开关解耦、轻度/严重超时并行报警、意图报警及质量闭环。来源链接：https://patents.google.com/patent/CN121814728A/zh

    #### 方向三：群聊意图识别、客服会话摘要和智能体协作

    6. **CN121480730A，面向群体会话的意图识别与任务自动派生方法**：该公开方案侧重基于因果网络的群体会话意图识别、任务派生和完整性验证，说明群体会话意图识别已有研究基础。其技术路线偏推理网络和任务派生，不公开本方案面向客户大群服务运营的切分、抑制、聚合、派单、报警和指标反馈闭环。来源链接：https://patents.google.com/patent/CN121480730A/zh

    7. **CN112686674A，客服会话工单摘要方法**：该方案按业务场景和主题边界拆解客服会话并输出工单摘要，说明“会话切分 + 工单摘要”已有公开。其局限在于偏离线摘要与槽位提取，不涉及实时派单、重复抑制、多人聚合、分级报警和运营指标闭环。来源链接：https://patents.google.com/patent/CN112686674A/zh

    8. **AliMe Assist 与 ISPY 等研究资料**：公开文献表明智能客服问答、多轮对话和群聊问题抽取属于已有研究方向。该类技术可作为背景，但通常不解决企业客户大群中的可派单会话任务生成、动态派单、分级报警和服务质量闭环。来源链接：https://arxiv.org/abs/1801.05032；https://arxiv.org/abs/2109.07055

    **相似技术写法参考结论**：同类正式文本通常先说明“公开方案能做什么”，再说明“未公开哪些组合特征”。因此，本稿不把“使用大模型识别意图”作为唯一创新点，而把可落地的会话窗口、抑制聚合、动态派单、分级监控和指标反馈写成连续技术链路。

    **本发明与现有技术的本质区别**：本发明不是单纯把群消息转为工单，而是以同群同客户为粒度生成会话任务，并在派单前后分别设置抑制、聚合、联合决策、分级报警和质量反馈机制；上述机制互相约束，形成从客户群消息到服务工单、从服务响应到策略优化的闭环。

    ### 1.2 现有技术存在的缺点

    1. **诉求上下文不足或被割裂**：逐条消息判断会导致单句语料不足，多句诉求又可能被拆成多个工单。
    2. **连续发言迟迟不派单**：仅依赖消息间隔时，客户持续补充信息会使会话长期不闭合。
    3. **多人交叉发言导致重复派单或误派**：同一事件可能被多个客户反复提及，帮扶者发言也可能被误判为新诉求。
    4. **派单对象缺少动态适配**：不同群类型、不同意图、不同时间段的处理人不同，固定名单或单一排班难以适配。
    5. **超时监控粒度不足**：单级提醒难以区分普通咨询、严重投诉和高风险意图，也难以支持“暂不派单但仍监控”的运营模式。
    6. **识别质量缺少闭环**：仅生成工单不能持续降低漏单和多派，无法根据人工复核结果反向优化阈值和兜底策略。

    ---

    ## 二、本发明所要解决的技术问题

    本发明要解决的技术问题是：在企业即时通信客户大群中，如何将多人交叉、多格式、持续变化的客户消息转化为可分派、可监控、可质检的服务工单，并在不显著增加漏单的情况下控制重复派单，进一步实现响应超时分级报警和识别质量闭环优化。

    具体包括：

    1. 通过同群同客户聚合和双阈值切分，解决上下文不足和连续发言不闭合的问题。
    2. 通过重复抑制和多人聚合，解决同一诉求重复派单、帮扶者发言误派和多人同事件多派的问题。
    3. 通过群类型、意图类别、时段和接收对象的联合决策，解决派单目标与业务角色、值班时段不匹配的问题。
    4. 通过轻度超时、严重超时和意图报警并行，解决普通超时和高风险意图无法差异化升级的问题。
    5. 通过人工复核与指标反馈，解决策略固定、模型漂移和业务变化导致的准确率下降问题。

    ---

    ## 三、本发明技术方案的详细阐述

    ### 3.1 背景

    本发明适用于企业即时通信客户群的服务运营场景。实施例可为售后服务客户群，包括区域群、门店群、一对一客户群等群类型，以及维修、保养、备件、权益、投诉等意图类别；上述行业和类别仅作为实施例，不作为权利要求限制。

    客户可通过文字、语音、图片、视频、链接、文件、小程序等方式表达诉求。系统通过托管账号、机器人或平台接口接入客户群消息，先生成会话任务，再识别意图并派单，同时对服务响应进行超时监控和人工复核。

    ### 3.2 系统框图

    ![系统框图](visio_figures/system_architecture.jpg)

    系统包括接入层、处理层、决策层、监控层和反馈层。接入层负责消息采集、身份归并和名单过滤；处理层负责会话聚合、双阈值任务生成、格式归一化和意图识别；决策层负责重复抑制、多人聚合和联合派单；监控层负责超时监听、分级报警和意图报警；反馈层负责人工复核、指标统计和策略调参。

    ### 3.3 模块功能说明

    1. **消息采集模块**：采集客户群消息事件，记录群标识、发言人标识、消息时间、消息格式和消息内容。
    2. **身份归并模块**：将同一群内同一客户的消息归入同一上下文，并区分客户、员工、机器人等身份。
    3. **过滤模块**：根据员工名单、机器人名单、客户黑名单过滤非客户诉求消息。
    4. **会话聚合模块**：按群标识 \(g\) 与客户标识 \(u\) 维护会话缓存。
    5. **双阈值会话任务生成模块**：根据相邻消息间隔阈值 \(T_{\mathrm{gap}}\) 与会话持续时长上限 \(T_{\mathrm{len}}\) 生成会话任务。
    6. **格式归一化模块**：将文字、语音、图片、视频、链接、文件、小程序等转换为统一任务表示；低业务价值格式可仅记录不派单。
    7. **意图识别模块**：识别会话任务的母意图和子意图，输出置信度与可派单类别；闲聊类不派单，其他类作为兜底类别。
    8. **联合派单决策模块**：依据群类型、意图类别、时段配置和接收对象确定派单目标。
    9. **重复抑制模块**：对同一客户同一事件在合并窗口 \(T_{\mathrm{merge}}\) 内的后续留言合并到已有工单。
    10. **多人聚合模块**：对多人无明确意图但持续讨论同一事件的会话任务进行聚合，汇总生成一个工单。
    11. **超时监听与报警模块**：对未有效回复的任务触发轻度超时、严重超时和特定意图报警。
    12. **反馈调参模块**：根据人工复核结果统计 \(Accuracy\)、\(MissRate\)、\(DupRate\)，并调整切分、抑制、聚合和兜底策略。

    ### 3.4 系统流程说明

    #### 流程图

    ![流程图](visio_figures/dispatch_monitor_flow.jpg)

    #### 流程说明

    1. 系统采集客户群消息事件，按照群标识与客户标识归并，并过滤员工、机器人和黑名单消息。
    2. 对同一群同一客户的连续消息维护会话缓存；当相邻消息间隔达到 \(T_{\mathrm{gap}}\) 或会话持续时长达到 \(T_{\mathrm{len}}\) 时，生成会话任务。
    3. 系统对会话任务进行格式归一化和意图识别，输出意图类别、置信度和可派单标记。
    4. 派单前先进行重复抑制；若同一客户同一或相似诉求在 \(T_{\mathrm{merge}}\) 内已有工单，则合并到已有工单。
    5. 若不命中重复抑制，则判断是否属于多人无明确意图但持续讨论同一事件；若满足聚合条件，则汇总生成一个多人聚合工单。
    6. 对普通会话任务，系统根据时段派送名单、意图处理人和默认处理人按优先级确定派单目标。
    7. 工单推送后建立超时监听；普通超时按轻度和严重阈值分级报警，高风险意图另行触发意图报警。
    8. 人工复核错派、漏派和多派结果，形成指标反馈并调整阈值和兜底策略。

    ### 3.4.1 符号与公式

    #### （1）符号与变量定义

    | 符号 | 含义 | 下标/量纲或取值范围 |
    |---|---|---|
    | \(g\) | 客户群标识 | 群级索引 |
    | \(u\) | 客户标识 | 客户级索引 |
    | \(k\) | 消息序号 | \(k=1,\ldots,n\) |
    | \(m_{g,u,k}\) | 客户 \(u\) 在群 \(g\) 中的第 \(k\) 条消息 | 消息对象 |
    | \(t_{g,u,k}\) | 消息 \(m_{g,u,k}\) 的发送时间 | 时间戳 |
    | \(s\) | 当前会话片段的起始消息序号 | \(1\leq s\leq k\) |
    | \(T_{\mathrm{gap}}\) | 相邻消息间隔阈值 | 秒，\(T_{\mathrm{gap}}>0\) |
    | \(T_{\mathrm{len}}\) | 会话持续时长上限 | 秒，\(T_{\mathrm{len}}>T_{\mathrm{gap}}\) 可选 |
    | \(x\) | 新生成的会话任务 | 任务对象 |
    | \(y\) | 已存在工单 | 工单对象 |
    | \(I(x)\) | 会话任务 \(x\) 的意图类别 | 离散类别 |
    | \(C(x)\) | 会话任务 \(x\) 的置信度 | \(0\leq C(x)\leq 1\) |
    | \(T_{\mathrm{merge}}\) | 重复抑制合并窗口 | 秒，\(T_{\mathrm{merge}}>0\) |
    | \(S(x,y)\) | 任务 \(x\) 与工单 \(y\) 的事件相似度 | \(0\leq S(x,y)\leq 1\) |
    | \(\theta\) | 事件相似度阈值 | \(0<\theta<1\) |
    | \(N_{\mathrm{user}}\) | 同一事件窗口内参与客户数 | 非负整数 |
    | \(A_{\mathrm{amb}}\) | 是否均为无明确意图任务 | 1 为是，0 为否 |
    | \(R(x)\) | 会话任务 \(x\) 是否已有有效回复 | 1 为已回复，0 为未回复 |
    | \(T_{\mathrm{light}}\) | 轻度超时阈值 | 秒，\(T_{\mathrm{light}}>0\) |
    | \(T_{\mathrm{severe}}\) | 严重超时阈值 | 秒，\(T_{\mathrm{severe}}>T_{\mathrm{light}}\) |
    | \(D_{\mathrm{slot}}\) | 当前时段配置的派送目标集合 | 可为空集合 |
    | \(D_{\mathrm{intent}}\) | 当前意图配置的处理目标集合 | 可为空集合 |
    | \(D_{\mathrm{default}}\) | 当前群类型对应的默认处理目标集合 | 可为空集合 |

    #### （2）核心公式

    **会话任务生成条件**：

    \[
    Close_{g,u,k}=1 \Longleftrightarrow \left(t_{g,u,k}-t_{g,u,k-1}\geq T_{\mathrm{gap}}\right)\lor\left(t_{g,u,k}-t_{g,u,s}\geq T_{\mathrm{len}}\right) \tag{1}
    \]

    当式 (1) 因第一项成立而触发时，系统闭合上一会话片段，并以当前消息作为下一会话片段的起点；当第二项成立时，系统闭合当前片段，避免持续发言导致任务长期不生成。

    **重复抑制判定条件**：

    \[
    Merge(x,y)=1 \Longleftrightarrow \left(g_x=g_y\right)\land\left(u_x=u_y\right)\land\left(I(x)=I(y)\lor S(x,y)\geq\theta\right)\land\left(t_x-t_y\leq T_{\mathrm{merge}}\right) \tag{2}
    \]

    式 (2) 表示：只有同一客户在同一群内、诉求相同或相似、且仍在合并窗口内时，才合并到已有工单。

    **多人聚合判定条件**：

    \[
    GroupTicket=1 \Longleftrightarrow \left(N_{\mathrm{user}}\geq 2\right)\land\left(A_{\mathrm{amb}}=1\right)\land\left(\overline{S}\geq\theta\right)\land\left(T_{\mathrm{duration}}\geq T_{\mathrm{len}}\right) \tag{3}
    \]

    其中 \(\overline{S}\) 为同一事件窗口内多个任务之间的平均相似度，\(T_{\mathrm{duration}}\) 为该事件窗口持续时长。

    **派单目标选择规则**：

    \[
    D(x)=\operatorname{first}\left(D_{\mathrm{slot}},D_{\mathrm{intent}},D_{\mathrm{default}}\right),\quad D(x)\neq\varnothing \tag{4}
    \]

    其中 \(\operatorname{first}(\cdot)\) 表示按顺序选择第一个非空目标集合；因此时段派送优先于意图处理人，意图处理人优先于默认处理人。

    **超时分级报警条件**：

    \[
    Severe(x)=1 \Longleftrightarrow R(x)=0\land now-t_{x,first}\geq T_{\mathrm{severe}} \tag{5}
    \]

    \[
    Light(x)=1 \Longleftrightarrow R(x)=0\land T_{\mathrm{light}}\leq now-t_{x,first}<T_{\mathrm{severe}} \tag{6}
    \]

    **质量闭环指标**：

    \[
    Accuracy=\frac{N_{\mathrm{correct}}}{N_{\mathrm{ticket}}},\quad MissRate=\frac{N_{\mathrm{miss}}}{N_{\mathrm{need}}},\quad DupRate=\frac{N_{\mathrm{dup}}}{N_{\mathrm{ticket}}} \tag{7}
    \]

    系统按照人工复核结果优先降低 \(MissRate\)，同时约束 \(DupRate\) 并保持 \(Accuracy\) 不低于预设要求。

    ### 3.5 关键技术参数

    | 符号 | 参数含义 | 取值范围/示例 | 说明 |
    |---|---|---|---|
    | \(T_{\mathrm{gap}}\) | 相邻消息间隔阈值 | 120-600 秒，示例 300 秒 | 触发会话切分 |
    | \(T_{\mathrm{len}}\) | 会话持续时长上限 | 300-900 秒，示例 600 秒 | 避免长期不派单 |
    | \(T_{\mathrm{merge}}\) | 重复抑制合并窗口 | 600-3600 秒，示例 1800 秒 | 同一诉求合并 |
    | \(\theta\) | 事件相似度阈值 | 0.5-0.9，示例 0.7 | 用于抑制和聚合 |
    | \(T_{\mathrm{light}}\) | 轻度超时阈值 | 300-1200 秒 | 触发轻度报警 |
    | \(T_{\mathrm{severe}}\) | 严重超时阈值 | 1200-3600 秒，且大于 \(T_{\mathrm{light}}\) | 触发严重报警 |
    | \(C(x)\) | 意图识别置信度 | 0-1 | 可用于低置信人工复核 |
    | \(Accuracy\) | 工单准确率 | 示例不低于 0.90 | 人工复核统计 |
    | \(MissRate\) | 漏单率 | 越低越好 | 优先控制指标 |
    | \(DupRate\) | 多派率 | 在不增加漏单前提下降低 | 抑制与聚合指标 |

    以上参数可根据业务、群规模和响应要求配置，不作为权利要求限制。

    ---

    ## 四、与现有技术相比，本发明具有的优点

    1. **会话任务生成更稳定**：双阈值切分同时解决“单句信息不足”和“连续发言不闭合”两个问题。
    2. **降低重复派单且兼顾漏单控制**：重复抑制处理同一客户同一诉求，多人聚合处理多人同事件讨论，使多派率受控而不简单牺牲漏单率。
    3. **派单策略适配业务角色和时间段**：时段派送、意图处理人和默认处理人具有明确优先级，适用于区域群、门店群、一对一群等不同群类型。
    4. **监控与派单解耦**：派单开关关闭时仍可生成工单或监控记录，便于质检、统计和后续复盘。
    5. **超时报警更精细**：轻度超时、严重超时和意图报警并行，可针对投诉抱怨等高风险意图提前升级。
    6. **可持续优化**：通过人工复核和三项质量指标形成反馈闭环，避免一次性规则或模型长期失效。
    7. **适配多格式客户表达**：对语音、图片、视频、链接、文件、小程序等进行格式归一化，降低非文本诉求漏识别风险。

    ---

    ## 五、本发明的技术关键点和欲保护点

    1. **同群同客户双阈值会话任务生成机制**：以群标识和客户标识为聚合键，以 \(T_{\mathrm{gap}}\) 和 \(T_{\mathrm{len}}\) 生成会话任务。
    2. **重复抑制与多人聚合组合机制**：在 \(T_{\mathrm{merge}}\) 内合并同一客户相同或相似诉求，并对多人无明确意图但同事件持续讨论的会话汇总派单。
    3. **群类型、意图、时段和目标对象联合派单机制**：按 \(D_{\mathrm{slot}}\)、\(D_{\mathrm{intent}}\)、\(D_{\mathrm{default}}\) 的优先级确定派单目标。
    4. **分级超时与意图报警并行机制**：基于 \(T_{\mathrm{light}}\)、\(T_{\mathrm{severe}}\) 触发轻度和严重报警，并对特定高风险意图触发独立意图报警。
    5. **派单开关与监控开关解耦机制**：支持不外部派单但仍生成工单记录、报警记录和质检数据。
    6. **多格式消息归一化与可派单意图识别机制**：将不同消息格式转换为统一任务表示，并结合意图体系决定是否派单。
    7. **基于 \(Accuracy\)、\(MissRate\)、\(DupRate\) 的质量闭环机制**：根据人工复核结果调整阈值、相似度、抑制窗口和兜底策略。

    ---

    ## 六、其它

    ### 6.1 实施例

    某企业将托管账号或机器人接入区域客户群、门店客户群和一对一客户群。系统监听群消息后，对客户消息进行身份归并和过滤；当同一客户连续发言达到双阈值条件时生成会话任务；系统对语音转写、图片和链接摘要等进行归一化，再识别维修、保养、备件、权益、投诉等意图。

    若客户在 30 分钟内持续补充同一诉求，系统将后续消息合并到已有工单；若多个客户围绕同一问题持续讨论且均无明确可派单意图，系统在达到持续时长和相似度条件后汇总派单。派单目标优先采用当前时段派送名单；未配置时段名单时，采用意图处理人；仍未配置时采用默认处理人。

    ### 6.2 技术效果

    在测试场景中，系统可将客户群中分散的消息转化为千级工单样本，工单准确率可达到 90% 以上，漏单率可控制为 0，多派率可控制在 30% 以下。上述结果说明，本方案能够提升客户响应时效、缩短信传递链路，并为客服效率提升和群内服务去人化提供技术基础。

    ### 6.3 可选实施方式

    1. 对低置信度意图任务设置人工复核队列，复核后再派单或调整目标。
    2. 对投诉抱怨、服务态度差、维修能力抱怨等高风险意图设置更短报警阈值。
    3. 对方言语音、车型简称、故障部件、谐音词和专业同义词设置领域词典或纠错模型。
    4. 对不同群类型设置不同 \(T_{\mathrm{gap}}\)、\(T_{\mathrm{len}}\)、\(T_{\mathrm{merge}}\) 和 \(\theta\)，以适配群规模和服务强度。
    5. 对派单结果进行可视化统计，形成错派、漏派、多派原因分布，用于策略更新。
    """
).strip()


PLAN = dedent(
    """
    # 项目可扩展专利方案清单

    ## 一、可继续产出的专利方案

    | 序号 | 方案名称建议 | 核心保护点 | 成熟度 | 建议优先级 |
    |---|---|---|---|---|
    | 1 | 一种基于同群同客户双阈值的会话任务生成与意图派单方法 | 群ID+客户ID聚合、间隔阈值+持续时长上限、意图派单 | 已在总案中充分体现，可拆分 | 高 |
    | 2 | 一种面向多人交叉客户群聊的诉求聚合与重复抑制方法 | 同一客户合并窗口、多人相似事件聚合、漏单/多派平衡 | 已有测试依据 | 高 |
    | 3 | 一种融合群类型、意图类别和时段配置的动态派单方法 | 时段派送优先、意图处理人、默认处理人、个人/群账号派送 | 已有产品配置依据 | 中高 |
    | 4 | 一种企业客户群响应超时分级报警与意图报警并行方法 | 轻度超时、严重超时、高风险意图报警、派单/监控解耦 | 已有产品配置依据 | 高 |
    | 5 | 一种面向售后客户群的多格式消息归一化与工单生成方法 | 语音转写、图片/视频/链接/文件摘要、低价值格式过滤 | 需补充技术细节 | 中高 |
    | 6 | 一种基于准确率、漏单率和多派率的派单策略闭环优化方法 | 人工复核、三指标统计、阈值与兜底策略调参 | 需要明确自动调参实现 | 中高 |
    | 7 | 一种针对高风险投诉意图的自动安抚与人工升级联动方法 | 投诉识别、安抚回复、主管/客服升级、短阈值监听 | 可作为后续增强 | 中 |
    | 8 | 一种结合领域词典的方言语音诉求识别与车型故障词纠错方法 | 语音转写纠错、车型/部件/故障词库、意图识别融合 | 需要研发补证 | 中 |
    | 9 | 一种客户群服务质检数据看板与策略推荐方法 | 响应时长、错派原因、漏派原因、多派原因、策略建议 | 需补充看板与推荐逻辑 | 中 |
    | 10 | 一种企业即时通信托管账号的群消息合规采集与脱敏处理方法 | 托管账号接入、消息脱敏、角色过滤、审计留痕 | 需确认合规实现 | 中 |

    ## 二、建议布局

    1. **总案**：保留当前“智能会话派单与运营监控”作为总案，覆盖采集、切分、识别、派单、报警和闭环。
    2. **第一批分案**：优先拆“多人交叉群聊的诉求聚合与重复抑制”“响应超时分级报警与意图报警并行”两案，区别点清楚且与公开工单技术差异大。
    3. **第二批增强案**：在补足研发细节后，再写“多格式消息归一化”“三指标闭环优化”“领域词典语音纠错”。
    4. **撰写注意**：不要把“大模型识别意图”单独作为发明点；应把模型输出与具体业务约束、阈值、派单对象、报警状态和复核指标结合起来写。
    """
).strip()


def add_md_artifacts():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    make_system_fig(FIG_DIR / "system_architecture.jpg")
    make_flow_fig(FIG_DIR / "dispatch_monitor_flow.jpg")
    MD_PATH.write_text(DISCLOSURE + "\n", encoding="utf-8")
    PLAN_PATH.write_text(PLAN + "\n", encoding="utf-8")


def set_run_font(run, size=10.5, bold=False, color=None, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if text.startswith("MATH:"):
        add_simple_math(p, text[5:], para=False)
    else:
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=bold)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell_text(table.cell(i, j), val, bold=(i == 0))
    doc.add_paragraph()


def add_para(doc, text="", size=10.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_run_font(r, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True, name="黑体", color="1F4D78")


def math_run(text: str):
    run = OxmlElement("m:r")
    rpr = OxmlElement("m:rPr")
    sty = OxmlElement("m:sty")
    sty.set(qn("m:val"), "p")
    rpr.append(sty)
    txt = OxmlElement("m:t")
    txt.text = text
    run.append(rpr)
    run.append(txt)
    return run


def math_sub(base: str, subscript: str):
    node = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(math_run(base))
    sub = OxmlElement("m:sub")
    sub.append(math_run(subscript))
    node.append(e)
    node.append(sub)
    return node


def math_frac(num, den):
    node = OxmlElement("m:f")
    n = OxmlElement("m:num")
    append_math_parts(n, num)
    d = OxmlElement("m:den")
    append_math_parts(d, den)
    node.append(n)
    node.append(d)
    return node


def append_math_parts(parent, parts):
    for part in parts:
        if isinstance(part, str):
            parent.append(math_run(part))
        elif part[0] == "sub":
            parent.append(math_sub(part[1], part[2]))
        elif part[0] == "frac":
            parent.append(math_frac(part[1], part[2]))


def add_math_parts(paragraph, parts, para: bool = True):
    math = OxmlElement("m:oMath")
    append_math_parts(math, parts)
    if para:
        math_para = OxmlElement("m:oMathPara")
        math_para.append(math)
        paragraph._p.append(math_para)
    else:
        paragraph._p.append(math)


def add_simple_math(paragraph, text: str, para: bool = True):
    parts = []
    i = 0
    while i < len(text):
        start = text.find("[", i)
        if start == -1:
            parts.append(text[i:])
            break
        base_start = start - 1
        while base_start > i and (text[base_start - 1].isalnum() or text[base_start - 1] == "_"):
            base_start -= 1
        if base_start > i:
            parts.append(text[i:base_start])
        end = text.find("]", start)
        if end == -1:
            parts.append(text[base_start:])
            break
        base = text[base_start:start]
        sub = text[start + 1:end]
        parts.append(("sub", base, sub))
        i = end + 1
    add_math_parts(paragraph, parts, para=para)


def add_formula(doc, parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_math_parts(p, parts, para=True)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    add_heading(doc, "技术交底书", 1)
    for line in [
        f"案件名称：{CASE}",
        "技术联系人：姓名/电话/邮箱待填写",
        "专利类型：发明",
    ]:
        add_para(doc, line, bold=True)

    add_heading(doc, "一、技术背景与现有技术", 1)
    add_heading(doc, "1.1 现有技术", 2)
    add_para(doc, "检索说明：本稿参考 CNIPA/Google Patents、企业即时通信及智能客服公开资料，以客户群消息、群聊工单、会话切分、意图识别、智能派单、超时报警和运营监控等为检索词进行比对。")
    add_table(doc, [
        ["方向", "代表公开资料", "与本方案的区别"],
        ["企业即时通信/SCRM", "企业微信、微伴助手、微盛企微管家、尘锋 SCRM", "公开能力偏通信、私域运营或通用工单，未见双阈值切分、抑制聚合、动态派单、分级报警和质量闭环组合。"],
        ["群聊消息到工单", "CN116032871B、CN122226741A、CN121814728A", "公开群消息建单或会话分配提醒，但未公开本方案的会话任务生成、重复抑制、多人聚合、时段优先派单和闭环。"],
        ["意图识别/摘要/智能体", "CN121480730A、CN112686674A、AliMe Assist、ISPY", "侧重意图识别、摘要或智能体协作，不构成客户大群派单与运营监控闭环。"],
    ])
    add_para(doc, "本发明与现有技术的本质区别在于：以同群同客户为粒度生成会话任务，并在派单前后设置抑制、聚合、联合决策、分级报警和质量反馈机制，形成从客户群消息到服务工单、从服务响应到策略优化的闭环。")
    add_heading(doc, "1.2 现有技术存在的缺点", 2)
    for item in [
        "诉求上下文不足或被割裂，逐条消息判断导致单句语料不足，多句诉求又可能被拆成多个工单。",
        "连续发言迟迟不派单，仅依赖消息间隔时会话可能长期不闭合。",
        "多人交叉发言导致重复派单或误派，同一事件可能被多个客户反复提及。",
        "派单对象缺少动态适配，固定名单或单一排班难以适配群类型、意图和时段差异。",
        "超时监控粒度不足，单级提醒难以区分普通咨询、严重投诉和高风险意图。",
        "识别质量缺少闭环，不能根据人工复核结果反向优化阈值和兜底策略。",
    ]:
        add_para(doc, item)

    add_heading(doc, "二、本发明所要解决的技术问题", 1)
    add_para(doc, "本发明要解决的技术问题是：在企业即时通信客户大群中，如何将多人交叉、多格式、持续变化的客户消息转化为可分派、可监控、可质检的服务工单，并实现响应超时分级报警和识别质量闭环优化。")

    add_heading(doc, "三、本发明技术方案的详细阐述", 1)
    add_heading(doc, "3.1 背景", 2)
    add_para(doc, "本发明适用于企业即时通信客户群的服务运营场景。客户可通过文字、语音、图片、视频、链接、文件、小程序等方式表达诉求。系统通过托管账号、机器人或平台接口接入客户群消息，先生成会话任务，再识别意图并派单，同时进行超时监控和人工复核。")
    add_heading(doc, "3.2 系统框图", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / "system_architecture.jpg"), width=Inches(6.4))
    add_para(doc, "系统包括接入层、处理层、决策层、监控层和反馈层；反馈层以人工复核和指标统计结果反向调整切分、抑制、聚合与兜底派单策略。")
    add_heading(doc, "3.3 模块功能说明", 2)
    for item in [
        "消息采集、身份归并与过滤模块负责获取客户群消息事件，区分客户、员工和机器人身份。",
        "会话聚合与双阈值任务生成模块按群标识与客户标识维护缓存，根据 T_gap（相邻消息间隔阈值）和 T_len（会话持续时长上限）生成会话任务。",
        "格式归一化与意图识别模块将不同消息格式转换为统一任务表示，并输出意图类别、置信度和可派单标记。",
        "重复抑制、多人聚合和联合派单模块在派单前处理同一诉求合并、多人同事件聚合，并按时段、意图和默认目标确定派单对象。",
        "超时监听、分级报警与意图报警模块对未有效回复任务触发轻度、严重和高风险意图报警。",
        "人工复核、指标统计与策略调参模块统计准确率、漏单率和多派率，形成闭环。",
    ]:
        add_para(doc, item)
    add_heading(doc, "3.4 系统流程说明", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG_DIR / "dispatch_monitor_flow.jpg"), width=Inches(6.4))
    add_para(doc, "流程为：采集群消息、身份归并与过滤、同群同客户会话聚合、双阈值切分、格式归一化、意图识别、重复抑制或多人聚合、联合派单、超时/意图报警、人工复核和策略调参。")

    add_heading(doc, "3.4.1 符号与公式", 2)
    add_table(doc, [
        ["符号", "含义", "下标/量纲或取值范围"],
        ["MATH:g, u, k", "群标识、客户标识、消息序号", "群级/客户级/序号索引"],
        ["MATH:m[g,u,k], t[g,u,k]", "客户 u 在群 g 的第 k 条消息及其发送时间", "消息对象/时间戳"],
        ["MATH:T[gap], T[len]", "相邻消息间隔阈值/会话持续时长上限", "秒"],
        ["MATH:x, y", "新会话任务/已有工单", "任务对象/工单对象"],
        ["MATH:I(x), C(x)", "任务 x 的意图类别和置信度", "离散类别/0-1"],
        ["MATH:T[merge]", "重复抑制合并窗口", "秒"],
        ["MATH:S(x,y), θ", "事件相似度及阈值", "0-1"],
        ["MATH:T[light], T[severe]", "轻度/严重超时阈值", "秒，T_severe>T_light"],
        ["MATH:D[slot], D[intent], D[default]", "时段、意图和默认派送目标集合", "可为空集合"],
    ])
    for f in [
        [("sub", "Close", "g,u,k"), "=1 ⇔ (", ("sub", "Δt", "k"), " ≥ ", ("sub", "T", "gap"), ") ∨ (", ("sub", "L", "g,u,k"), " ≥ ", ("sub", "T", "len"), ")    (1)"],
        ["Merge(x,y)=1 ⇔ (", ("sub", "g", "x"), "=", ("sub", "g", "y"), ") ∧ (", ("sub", "u", "x"), "=", ("sub", "u", "y"), ") ∧ (I(x)=I(y) ∨ S(x,y)≥θ) ∧ (", ("sub", "Δt", "x,y"), "≤", ("sub", "T", "merge"), ")    (2)"],
        ["GroupTicket=1 ⇔ (", ("sub", "N", "user"), "≥2) ∧ (", ("sub", "A", "amb"), "=1) ∧ (", ("sub", "S", "avg"), "≥θ) ∧ (", ("sub", "T", "duration"), "≥", ("sub", "T", "len"), ")    (3)"],
        ["D(x)=first(", ("sub", "D", "slot"), ", ", ("sub", "D", "intent"), ", ", ("sub", "D", "default"), "),  D(x)≠∅    (4)"],
        ["Severe(x)=1 ⇔ R(x)=0 ∧ now-", ("sub", "t", "x,first"), "≥", ("sub", "T", "severe"), "    (5)"],
        ["Light(x)=1 ⇔ R(x)=0 ∧ ", ("sub", "T", "light"), "≤now-", ("sub", "t", "x,first"), "<", ("sub", "T", "severe"), "    (6)"],
        ["Accuracy=", ("frac", [("sub", "N", "correct")], [("sub", "N", "ticket")]), ",  MissRate=", ("frac", [("sub", "N", "miss")], [("sub", "N", "need")]), ",  DupRate=", ("frac", [("sub", "N", "dup")], [("sub", "N", "ticket")]), "    (7)"],
    ]:
        add_formula(doc, f)
    add_heading(doc, "3.5 关键技术参数", 2)
    add_table(doc, [
        ["符号", "参数含义", "取值范围/示例", "说明"],
        ["MATH:T[gap]", "相邻消息间隔阈值", "120-600 秒，示例 300 秒", "触发会话切分"],
        ["MATH:T[len]", "会话持续时长上限", "300-900 秒，示例 600 秒", "避免长期不派单"],
        ["MATH:T[merge]", "重复抑制合并窗口", "600-3600 秒，示例 1800 秒", "同一诉求合并"],
        ["MATH:θ", "事件相似度阈值", "0.5-0.9，示例 0.7", "抑制和聚合共用"],
        ["MATH:T[light]", "轻度超时阈值", "300-1200 秒", "轻度报警"],
        ["MATH:T[severe]", "严重超时阈值", "1200-3600 秒", "严重报警"],
    ])

    add_heading(doc, "四、与现有技术相比的优点", 1)
    for item in [
        "会话任务生成更稳定，双阈值切分同时解决单句信息不足和连续发言不闭合。",
        "重复抑制与多人聚合降低重复派单，并兼顾漏单控制。",
        "派单策略适配业务角色和时间段，时段派送、意图处理人和默认处理人具有明确优先级。",
        "监控与派单解耦，便于质检、统计和后续复盘。",
        "超时报警更精细，轻度超时、严重超时和高风险意图报警并行。",
        "通过人工复核和三项质量指标形成可持续优化闭环。",
    ]:
        add_para(doc, item)
    add_heading(doc, "五、本发明的技术关键点和欲保护点", 1)
    for item in [
        "同群同客户双阈值会话任务生成机制。",
        "重复抑制与多人聚合组合机制。",
        "群类型、意图、时段和目标对象联合派单机制。",
        "分级超时与意图报警并行机制。",
        "派单开关与监控开关解耦机制。",
        "多格式消息归一化与可派单意图识别机制。",
        "基于 Accuracy、MissRate、DupRate 的质量闭环机制。",
    ]:
        add_para(doc, item)
    add_heading(doc, "六、其它", 1)
    add_para(doc, "实施例：某企业将托管账号或机器人接入区域客户群、门店客户群和一对一客户群。系统监听群消息后，对客户消息进行身份归并和过滤；当同一客户连续发言达到双阈值条件时生成会话任务；系统对语音转写、图片和链接摘要等进行归一化，再识别维修、保养、备件、权益、投诉等意图。")
    add_para(doc, "技术效果：在测试场景中，系统可将客户群中分散消息转化为千级工单样本，工单准确率可达到 90% 以上，漏单率可控制为 0，多派率可控制在 30% 以下。")

    doc.save(DOCX_PATH)


def append_log():
    local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    utc = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    entry = f"""

    ## {local} / {utc}

    - 类型：合并迭代
    - 用户说明摘要：将公式直接插入 Word，公式与符号采用论文式下标和分式规范；图内不写图标题，流程连线清楚且不交叉。
    - 本轮交付文件：{MD_PATH.name}；{DOCX_PATH.name}；{PLAN_PATH.name}
    - 合并摘要摘录：将 Word 公式由普通文本改为原生 OMML 数学对象，统一 3.4.1 与 3.5 的符号下标体例；重绘系统框图和流程图，去掉图内标题并消除交叉连线。
    """
    LOG_PATH.write_text((LOG_PATH.read_text(encoding="utf-8") if LOG_PATH.exists() else "# 交底书修订对话记录\n") + dedent(entry), encoding="utf-8")


def main():
    add_md_artifacts()
    build_docx()
    append_log()
    print(MD_PATH)
    print(DOCX_PATH)
    print(PLAN_PATH)


if __name__ == "__main__":
    main()
